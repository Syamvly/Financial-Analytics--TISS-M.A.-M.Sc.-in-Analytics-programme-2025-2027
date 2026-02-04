import docx
import sys

def convert_docx_to_md(docx_path, md_path):
    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        print(f"Error loading docx: {e}")
        return

    markdown_lines = []
    
    # We want to iterate in document order. 
    # doc.paragraphs only gives paragraphs (outside tables).
    # doc.tables only gives tables.
    # To get them in order, we inspect the XML structure or use iter_block_items logic if available,
    # but a simpler robust way for simple docs is often iterating body elements.
    
    def iter_block_items(parent):
        """
        Yield each paragraph and table child within parent, in document order.
        Each returned value is an instance of either Table or Paragraph. 
        """
        if isinstance(parent, docx.document.Document):
            parent_elm = parent.element.body
        elif isinstance(parent, docx.table._Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, docx.oxml.text.paragraph.CT_P):
                yield docx.text.paragraph.Paragraph(child, parent)
            elif isinstance(child, docx.oxml.table.CT_Tbl):
                yield docx.table.Table(child, parent)

    for block in iter_block_items(doc):
        if isinstance(block, docx.text.paragraph.Paragraph):
            text = block.text.strip()
            if text:
                # Basic header detection (heuristic based on style or length/content)
                if block.style.name.startswith('Heading 1'):
                    markdown_lines.append(f"# {text}")
                elif block.style.name.startswith('Heading 2'):
                    markdown_lines.append(f"## {text}")
                elif block.style.name.startswith('Heading 3'):
                    markdown_lines.append(f"### {text}")
                else:
                    markdown_lines.append(text)
                markdown_lines.append("") # Add spacing
                
        elif isinstance(block, docx.table.Table):
            # Convert table to markdown
            rows = block.rows
            if not rows:
                continue
                
            # Extract distinct text from first row to use as header or just first row
            header_cells = [cell.text.strip().replace('\n', ' ') for cell in rows[0].cells]
            
            # Markdown table header
            markdown_lines.append("| " + " | ".join(header_cells) + " |")
            markdown_lines.append("| " + " | ".join(['---'] * len(header_cells)) + " |")
            
            # Data rows
            for row in rows[1:]:
                row_cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                markdown_lines.append("| " + " | ".join(row_cells) + " |")
            
            markdown_lines.append("") # Spacing

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(markdown_lines))
    
    print(f"Successfully converted {docx_path} to {md_path}")

if __name__ == "__main__":
    docx_file = "FA Course Syllabus & Structure.docx"
    md_file = "README.md"
    convert_docx_to_md(docx_file, md_file)
